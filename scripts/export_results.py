#!/usr/bin/env python3
"""
Script de Exportación de Resultados
====================================

Este script permite exportar notebooks de Jupyter a PDF, copiar gráficas
generadas, crear reportes en formato Word y generar archivos ZIP con todos
los resultados del proyecto TIF Cálculo Fase III.

Autor: Aron
Universidad: UCSM
Curso: Cálculo 2025 - Fase III

Uso:
    python3 scripts/export_results.py
    python3 scripts/export_results.py --format pdf
    python3 scripts/export_results.py --format all --zip

Requisitos:
    - nbconvert: Para convertir notebooks a PDF
    - python-docx: Para crear documentos Word
    - Jupyter instalado y configurado
"""

import os
import sys
import shutil
import zipfile
import argparse
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional

# Intentar importar dependencias opcionales
try:
    import nbformat
    from nbconvert import PDFExporter, HTMLExporter
    from nbconvert.preprocessors import ExecutePreprocessor
    NBCONVERT_AVAILABLE = True
except ImportError:
    NBCONVERT_AVAILABLE = False
    print("⚠️  nbconvert no está instalado. Instalarlo con: pip install nbconvert")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx no está instalado. Instalarlo con: pip install python-docx")


class ResultsExporter:
    """
    Clase principal para exportar resultados del proyecto TIF Cálculo Fase III.
    """

    def __init__(self, project_root: str = None):
        """
        Inicializa el exportador de resultados.

        Args:
            project_root: Ruta raíz del proyecto. Si es None, se detecta automáticamente.
        """
        if project_root is None:
            # Detectar automáticamente la raíz del proyecto
            script_dir = Path(__file__).resolve().parent
            self.project_root = script_dir.parent
        else:
            self.project_root = Path(project_root)

        # Definir rutas importantes
        self.notebooks_dir = self.project_root / "services" / "jupyter" / "notebooks"
        self.plots_dir = self.project_root / "shared" / "plots"
        self.results_dir = self.project_root / "shared" / "results"
        self.animations_dir = self.project_root / "shared" / "animations"
        self.exports_dir = self.results_dir / "exports"

        # Crear directorios si no existen
        self.exports_dir.mkdir(parents=True, exist_ok=True)

        # Timestamp para nombres de archivo
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def find_notebooks(self) -> List[Path]:
        """
        Encuentra todos los notebooks en el directorio de Jupyter.

        Returns:
            Lista de rutas a archivos .ipynb
        """
        if not self.notebooks_dir.exists():
            print(f"❌ Directorio de notebooks no encontrado: {self.notebooks_dir}")
            return []

        notebooks = list(self.notebooks_dir.glob("*.ipynb"))

        # Filtrar checkpoints de Jupyter
        notebooks = [nb for nb in notebooks if ".ipynb_checkpoints" not in str(nb)]

        return sorted(notebooks)

    def export_notebook_to_pdf(self, notebook_path: Path) -> Optional[Path]:
        """
        Exporta un notebook de Jupyter a PDF usando nbconvert.

        Args:
            notebook_path: Ruta al archivo .ipynb

        Returns:
            Ruta al archivo PDF generado, o None si falla
        """
        if not NBCONVERT_AVAILABLE:
            print(f"⚠️  No se puede exportar {notebook_path.name}: nbconvert no disponible")
            return None

        try:
            print(f"📄 Exportando {notebook_path.name} a PDF...")

            # Leer el notebook
            with open(notebook_path, 'r', encoding='utf-8') as f:
                nb = nbformat.read(f, as_version=4)

            # Configurar el exportador PDF
            pdf_exporter = PDFExporter()
            pdf_exporter.template_name = 'classic'

            # Exportar a PDF
            (body, resources) = pdf_exporter.from_notebook_node(nb)

            # Guardar el PDF
            pdf_filename = f"{notebook_path.stem}_{self.timestamp}.pdf"
            pdf_path = self.exports_dir / pdf_filename

            with open(pdf_path, 'wb') as f:
                f.write(body)

            print(f"   ✓ PDF generado: {pdf_path}")
            return pdf_path

        except Exception as e:
            print(f"   ❌ Error al exportar {notebook_path.name}: {e}")
            return None

    def export_notebook_to_html(self, notebook_path: Path) -> Optional[Path]:
        """
        Exporta un notebook de Jupyter a HTML usando nbconvert.

        Args:
            notebook_path: Ruta al archivo .ipynb

        Returns:
            Ruta al archivo HTML generado, o None si falla
        """
        if not NBCONVERT_AVAILABLE:
            print(f"⚠️  No se puede exportar {notebook_path.name}: nbconvert no disponible")
            return None

        try:
            print(f"🌐 Exportando {notebook_path.name} a HTML...")

            # Leer el notebook
            with open(notebook_path, 'r', encoding='utf-8') as f:
                nb = nbformat.read(f, as_version=4)

            # Configurar el exportador HTML
            html_exporter = HTMLExporter()
            html_exporter.template_name = 'classic'

            # Exportar a HTML
            (body, resources) = html_exporter.from_notebook_node(nb)

            # Guardar el HTML
            html_filename = f"{notebook_path.stem}_{self.timestamp}.html"
            html_path = self.exports_dir / html_filename

            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(body)

            print(f"   ✓ HTML generado: {html_path}")
            return html_path

        except Exception as e:
            print(f"   ❌ Error al exportar {notebook_path.name}: {e}")
            return None

    def copy_plots(self) -> List[Path]:
        """
        Copia todas las gráficas del directorio plots/ a results/exports/.

        Returns:
            Lista de archivos copiados
        """
        if not self.plots_dir.exists():
            print(f"⚠️  Directorio de gráficas no encontrado: {self.plots_dir}")
            return []

        copied_files = []

        # Extensiones de imagen soportadas
        image_extensions = ['.png', '.jpg', '.jpeg', '.svg', '.pdf']

        print(f"\n📊 Copiando gráficas desde {self.plots_dir}...")

        for img_file in self.plots_dir.iterdir():
            if img_file.suffix.lower() in image_extensions:
                dest_file = self.exports_dir / img_file.name
                shutil.copy2(img_file, dest_file)
                copied_files.append(dest_file)
                print(f"   ✓ Copiado: {img_file.name}")

        if copied_files:
            print(f"   Total: {len(copied_files)} archivo(s) copiado(s)")
        else:
            print(f"   ⚠️  No se encontraron gráficas para copiar")

        return copied_files

    def create_word_report(self, notebooks: List[Path], plots: List[Path]) -> Optional[Path]:
        """
        Crea un reporte en formato Word con información del proyecto.

        Args:
            notebooks: Lista de notebooks procesados
            plots: Lista de gráficas incluidas

        Returns:
            Ruta al archivo Word generado, o None si falla
        """
        if not DOCX_AVAILABLE:
            print("⚠️  No se puede crear reporte Word: python-docx no disponible")
            return None

        try:
            print(f"\n📝 Creando reporte en Word...")

            # Crear documento
            doc = Document()

            # Configurar estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            # Portada
            title = doc.add_heading('TIF Cálculo Fase III', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_heading('Aplicaciones de la Derivada', 2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Información del autor
            doc.add_paragraph()
            author_info = doc.add_paragraph()
            author_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_info.add_run('Autor: Aron\n').bold = True
            author_info.add_run('Universidad Católica de Santa María (UCSM)\n')
            author_info.add_run('Curso: Cálculo 2025 - Fase III\n')
            author_info.add_run(f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M")}\n')

            doc.add_page_break()

            # Resumen ejecutivo
            doc.add_heading('1. Resumen Ejecutivo', 1)
            doc.add_paragraph(
                'Este documento contiene los resultados del Trabajo de Investigación Formativa (TIF) '
                'sobre Aplicaciones de la Derivada, desarrollado como parte del curso de Cálculo Fase III. '
                'El proyecto implementa una plataforma computacional multi-motor para el análisis matemático '
                'utilizando software libre.'
            )

            # Objetivos
            doc.add_heading('2. Objetivos del Proyecto', 1)
            objectives = doc.add_paragraph()
            objectives.add_run('Objetivos principales:\n').bold = True
            doc.add_paragraph('• Implementar análisis de máximos y mínimos de funciones', style='List Bullet')
            doc.add_paragraph('• Estudiar concavidad y puntos de inflexión', style='List Bullet')
            doc.add_paragraph('• Desarrollar técnicas de trazo de curvas', style='List Bullet')
            doc.add_paragraph('• Comparar resultados entre diferentes motores computacionales', style='List Bullet')

            # Tecnologías
            doc.add_heading('3. Tecnologías Utilizadas', 1)
            doc.add_paragraph(
                'El proyecto utiliza una arquitectura basada en Docker con los siguientes componentes:'
            )
            doc.add_paragraph('• Python 3.11 con SymPy (cálculo simbólico)', style='List Bullet')
            doc.add_paragraph('• Jupyter Lab (notebooks interactivos)', style='List Bullet')
            doc.add_paragraph('• SageMath (álgebra computacional avanzada)', style='List Bullet')
            doc.add_paragraph('• GNU Octave (computación numérica)', style='List Bullet')
            doc.add_paragraph('• Streamlit (dashboard web)', style='List Bullet')

            # Notebooks procesados
            doc.add_heading('4. Notebooks Procesados', 1)
            if notebooks:
                doc.add_paragraph(
                    f'Se han procesado {len(notebooks)} notebook(s) en este reporte:'
                )
                for nb in notebooks:
                    doc.add_paragraph(f'• {nb.stem}', style='List Bullet')
            else:
                doc.add_paragraph('No se procesaron notebooks en esta ejecución.')

            # Gráficas incluidas
            doc.add_heading('5. Gráficas Generadas', 1)
            if plots:
                doc.add_paragraph(
                    f'Se han generado {len(plots)} gráfica(s) durante el análisis:'
                )
                for plot in plots[:10]:  # Limitar a 10 para no saturar el documento
                    doc.add_paragraph(f'• {plot.name}', style='List Bullet')
                if len(plots) > 10:
                    doc.add_paragraph(f'... y {len(plots) - 10} gráfica(s) adicional(es)')
            else:
                doc.add_paragraph('No se encontraron gráficas en este reporte.')

            # Resultados
            doc.add_heading('6. Resultados y Conclusiones', 1)
            doc.add_paragraph(
                'Los resultados del análisis demuestran la efectividad de los métodos del cálculo '
                'diferencial para determinar características importantes de funciones, incluyendo '
                'valores máximos y mínimos, puntos de inflexión y comportamiento general de las curvas.'
            )

            # Conclusiones
            doc.add_paragraph()
            doc.add_paragraph('Conclusiones principales:', style='List Bullet')
            doc.add_paragraph(
                '• Las herramientas de cálculo simbólico permiten automatizar el análisis de derivadas',
                style='List Bullet'
            )
            doc.add_paragraph(
                '• La visualización interactiva facilita la comprensión de conceptos matemáticos',
                style='List Bullet'
            )
            doc.add_paragraph(
                '• El software libre es una alternativa viable para análisis matemático avanzado',
                style='List Bullet'
            )

            # Referencias
            doc.add_page_break()
            doc.add_heading('7. Referencias Bibliográficas', 1)
            doc.add_paragraph(
                '• Stewart, J. (2012). Cálculo de una variable: Trascendentes tempranas (7ª ed.). '
                'Cengage Learning.'
            )
            doc.add_paragraph(
                '• Larson, R., & Edwards, B. (2016). Cálculo (10ª ed.). Cengage Learning.'
            )
            doc.add_paragraph(
                '• SymPy Development Team. (2024). SymPy Documentation. https://docs.sympy.org'
            )

            # Guardar documento
            report_filename = f"reporte_tif_calculo_{self.timestamp}.docx"
            report_path = self.exports_dir / report_filename
            doc.save(report_path)

            print(f"   ✓ Reporte Word generado: {report_path}")
            return report_path

        except Exception as e:
            print(f"   ❌ Error al crear reporte Word: {e}")
            return None

    def create_zip_archive(self, files: List[Path]) -> Optional[Path]:
        """
        Crea un archivo ZIP con todos los resultados exportados.

        Args:
            files: Lista de archivos a incluir en el ZIP

        Returns:
            Ruta al archivo ZIP generado, o None si falla
        """
        if not files:
            print("⚠️  No hay archivos para incluir en el ZIP")
            return None

        try:
            print(f"\n📦 Creando archivo ZIP con resultados...")

            # Nombre del archivo ZIP
            zip_filename = f"tif_calculo_resultados_{self.timestamp}.zip"
            zip_path = self.results_dir / zip_filename

            # Crear el archivo ZIP
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file in files:
                    if file.exists():
                        # Agregar archivo con ruta relativa
                        arcname = file.name
                        zipf.write(file, arcname)
                        print(f"   ✓ Agregado: {file.name}")

            file_size_mb = zip_path.stat().st_size / (1024 * 1024)
            print(f"\n   ✓ ZIP generado: {zip_path}")
            print(f"   📊 Tamaño: {file_size_mb:.2f} MB")
            print(f"   📁 Archivos incluidos: {len(files)}")

            return zip_path

        except Exception as e:
            print(f"   ❌ Error al crear archivo ZIP: {e}")
            return None

    def export_all(self, export_format: str = 'all', create_zip: bool = False) -> Dict[str, List[Path]]:
        """
        Ejecuta el proceso completo de exportación.

        Args:
            export_format: Formato de exportación ('pdf', 'html', 'all')
            create_zip: Si True, crea un archivo ZIP con todos los resultados

        Returns:
            Diccionario con las rutas de archivos generados por tipo
        """
        results = {
            'pdfs': [],
            'htmls': [],
            'plots': [],
            'reports': [],
            'zips': []
        }

        print("=" * 70)
        print("🚀 EXPORTADOR DE RESULTADOS - TIF CÁLCULO FASE III")
        print("=" * 70)
        print(f"📂 Proyecto: {self.project_root}")
        print(f"📅 Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
        print(f"📋 Formato: {export_format.upper()}")
        print("=" * 70)

        # Encontrar notebooks
        notebooks = self.find_notebooks()

        if not notebooks:
            print("\n⚠️  No se encontraron notebooks para exportar")
        else:
            print(f"\n📚 Notebooks encontrados: {len(notebooks)}")
            for nb in notebooks:
                print(f"   • {nb.name}")

            # Exportar notebooks
            print(f"\n{'='*70}")
            print("EXPORTANDO NOTEBOOKS")
            print("=" * 70)

            for notebook in notebooks:
                if export_format in ['pdf', 'all']:
                    pdf_path = self.export_notebook_to_pdf(notebook)
                    if pdf_path:
                        results['pdfs'].append(pdf_path)

                if export_format in ['html', 'all']:
                    html_path = self.export_notebook_to_html(notebook)
                    if html_path:
                        results['htmls'].append(html_path)

        # Copiar gráficas
        print(f"\n{'='*70}")
        print("COPIANDO GRÁFICAS")
        print("=" * 70)
        plots = self.copy_plots()
        results['plots'] = plots

        # Crear reporte Word
        print(f"\n{'='*70}")
        print("GENERANDO REPORTE WORD")
        print("=" * 70)
        report = self.create_word_report(notebooks, plots)
        if report:
            results['reports'].append(report)

        # Crear ZIP si se solicita
        if create_zip:
            print(f"\n{'='*70}")
            print("CREANDO ARCHIVO ZIP")
            print("=" * 70)

            all_files = (
                results['pdfs'] +
                results['htmls'] +
                results['plots'] +
                results['reports']
            )

            zip_path = self.create_zip_archive(all_files)
            if zip_path:
                results['zips'].append(zip_path)

        # Resumen final
        print(f"\n{'='*70}")
        print("✅ RESUMEN DE EXPORTACIÓN")
        print("=" * 70)
        print(f"📄 PDFs generados: {len(results['pdfs'])}")
        print(f"🌐 HTMLs generados: {len(results['htmls'])}")
        print(f"📊 Gráficas copiadas: {len(results['plots'])}")
        print(f"📝 Reportes Word: {len(results['reports'])}")
        print(f"📦 Archivos ZIP: {len(results['zips'])}")
        print(f"\n📁 Directorio de salida: {self.exports_dir}")
        print("=" * 70)

        return results


def main():
    """
    Función principal del script.
    """
    # Configurar argumentos de línea de comandos
    parser = argparse.ArgumentParser(
        description='Exporta resultados del proyecto TIF Cálculo Fase III',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ejemplos de uso:
  python3 scripts/export_results.py                    # Exportar todo (PDF + HTML)
  python3 scripts/export_results.py --format pdf       # Solo PDF
  python3 scripts/export_results.py --format html      # Solo HTML
  python3 scripts/export_results.py --zip              # Exportar y crear ZIP
  python3 scripts/export_results.py -f all -z          # Todo + ZIP (forma corta)
        """
    )

    parser.add_argument(
        '-f', '--format',
        choices=['pdf', 'html', 'all'],
        default='all',
        help='Formato de exportación de notebooks (default: all)'
    )

    parser.add_argument(
        '-z', '--zip',
        action='store_true',
        help='Crear archivo ZIP con todos los resultados'
    )

    parser.add_argument(
        '-p', '--project-root',
        type=str,
        default=None,
        help='Ruta raíz del proyecto (default: auto-detectar)'
    )

    args = parser.parse_args()

    try:
        # Crear exportador
        exporter = ResultsExporter(project_root=args.project_root)

        # Ejecutar exportación
        results = exporter.export_all(
            export_format=args.format,
            create_zip=args.zip
        )

        # Código de salida basado en resultados
        total_files = sum(len(files) for files in results.values())

        if total_files == 0:
            print("\n⚠️  No se generaron archivos")
            return 1
        else:
            print(f"\n✅ Exportación completada exitosamente")
            return 0

    except KeyboardInterrupt:
        print("\n\n⚠️  Exportación cancelada por el usuario")
        return 130
    except Exception as e:
        print(f"\n❌ Error durante la exportación: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
